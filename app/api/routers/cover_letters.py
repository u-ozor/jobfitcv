# app/api/routers/cover_letters.py

import os
import json
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.orm import Session
from pydantic import BaseModel

from app.database.models import Job, Variant
from app.database.session import get_db
from app.core.llm_client import generate_completion

router = APIRouter()

PROFILE_PATH  = "data/user_profile.json"
CHUNKS_PATH   = "data/resume_chunks.json"
OUTPUT_DIR    = "outputs/cover_letters"
DEFAULT_MODEL = "claude-sonnet-4-6"

# CL_PROVIDER is independent from the rewrite pipeline's LLM_PROVIDER so toggling
# one doesn't silently change the other. Defaults preserve current behavior exactly
# (Anthropic, claude-sonnet-4-6) — set CL_PROVIDER=ollama in .env to switch.
CL_PROVIDER = os.environ.get("CL_PROVIDER", "anthropic")

VALID_TONES = {"professional", "direct", "warm"}

TONE_CLOSING = {
    "professional": "Sincerely",
    "direct":       "Regards",
    "warm":         "Warmly",
}

SYSTEM_PROMPT = """\
You are a professional cover letter writer helping a job seeker land interviews.

Write cover letters that:
- Open by naming the candidate's single most relevant qualification mapped directly to THIS role and company — not a generic enthusiasm statement
- Lift specific language from the job description: the company's stated mission or values, named platforms or tools, the team name, specific responsibilities
- Map the candidate's concrete experience — numbers, tenure, named tools — to those specific JD requirements, not generically
- Include one distinctive phrase or observation that reveals personality without being a cliché ("passion for", "team player", "excited to apply")
- Close by naming the specific team or role and one outcome the company cares about based on the JD
- Are exactly 3 paragraphs: hook → experience mapping → closing ask
- Use first-person active voice; under 320 words total

Return only the 3 body paragraphs. No salutation, no closing line, no date, no address blocks.\
"""

TONE_ADDENDUM = {
    "professional": "",
    "direct":       " Lead with metrics and impact. Skip pleasantries. Be blunt about fit.",
    "warm":         " Let genuine interest in the company's mission come through. Use the company's own language about their purpose.",
}


class CoverLetterRequest(BaseModel):
    tone:     Optional[str] = "professional"
    briefing: Optional[str] = None


def _load_json(path):
    if not os.path.exists(path):
        return None
    with open(path) as f:
        return json.load(f)


def _build_experience_summary(chunks):
    if not chunks or not isinstance(chunks, list):
        return "(no resume chunks found)"
    lines = []
    for chunk in chunks:
        content = chunk.get("content", "")
        if not content:
            continue
        if chunk.get("type") in ("experience", "project"):
            lines.append(content)
        elif chunk.get("type") == "summary" and chunk.get("track") is None:
            lines.insert(0, content)
    return "\n".join(f"- {b}" for b in lines[:14]) or "(no bullets found)"


@router.post("/{job_id}/cover_letter")
def generate_cover_letter(
    job_id: str,
    body: CoverLetterRequest,
    db: Session = Depends(get_db)
):
    tone = (body.tone or "professional").lower()
    if tone not in VALID_TONES:
        raise HTTPException(status_code=400, detail=f"tone must be one of: {', '.join(VALID_TONES)}")

    job = db.query(Job).filter(Job.id == job_id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    if CL_PROVIDER == "anthropic":
        try:
            import anthropic  # noqa: F401
        except ImportError:
            raise HTTPException(status_code=500, detail="anthropic package not installed on server")
        if not os.environ.get("ANTHROPIC_API_KEY"):
            raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY not set — add it to .env and restart")

    profile = _load_json(PROFILE_PATH) or {}
    chunks  = _load_json(CHUNKS_PATH)
    experience_summary = _build_experience_summary(chunks)

    # Personal context — bio/summary, social links
    bio_lines = []
    if profile.get("cover_letter"):
        bio_lines.append(f"Bio/summary: {profile['cover_letter']}")
    if profile.get("linkedin"):
        bio_lines.append(f"LinkedIn: {profile['linkedin']}")
    if profile.get("github"):
        bio_lines.append(f"GitHub: {profile['github']}")
    personal_context = "\n".join(bio_lines) if bio_lines else "(not provided)"

    prompt = f"""\
Write a {tone} cover letter for this application.{TONE_ADDENDUM[tone]}

<candidate>
Name: {profile.get("name", "the candidate")}
Location: {profile.get("location", "")}
{personal_context}
</candidate>

<candidate_experience>
{experience_summary}
</candidate_experience>

<role>
Title: {job.title or "Unknown role"}
Company: {job.company or "the company"}
</role>

<job_description>
{(job.raw_text or "")[:4000]}
</job_description>

Before writing: identify the company's stated mission or values in the JD, the 2-3 most specific requirements, and any named platforms, tools, or team context. Use that exact language in the letter — not paraphrases.
Output ONLY the cover letter body. No subject line, no date, no preamble."""

    if body.briefing:
        prompt += f"\n\n<briefing>\n{body.briefing}\n</briefing>\nPrioritise the angles and framing in the briefing above when structuring the letter."

    model = os.environ.get("CL_MODEL", DEFAULT_MODEL)
    body  = generate_completion(
        prompt,
        system=SYSTEM_PROMPT,
        provider=CL_PROVIDER,
        model=model,
        max_tokens=600
    )

    name    = profile.get("name", "").strip()
    closing = TONE_CLOSING.get(tone, "Sincerely")
    sign_off = f"{closing},\n{name}" if name else f"{closing},"
    text = f"Dear Hiring Manager,\n\n{body}\n\n{sign_off}"

    out_dir  = os.path.join(OUTPUT_DIR, job_id)
    os.makedirs(out_dir, exist_ok=True)
    md_path  = os.path.join(out_dir, f"cover_letter_{tone}.md")
    txt_path = os.path.join(out_dir, f"cover_letter_{tone}.txt")

    with open(md_path, "w") as f:
        f.write(f"# Cover Letter — {job.title} @ {job.company}\n\n")
        f.write(f"*Tone: {tone} | Model: {model}*\n\n---\n\n")
        f.write(text + "\n")

    with open(txt_path, "w") as f:
        f.write(text + "\n")

    return {
        "job_id":    job_id,
        "tone":      tone,
        "text":      text,
        "md_path":   md_path,
        "txt_path":  txt_path
    }


@router.get("/{job_id}/cover_letter")
def get_cover_letter(job_id: str, tone: str = "professional"):
    tone = tone.lower()
    txt_path = os.path.join(OUTPUT_DIR, job_id, f"cover_letter_{tone}.txt")
    if not os.path.exists(txt_path):
        raise HTTPException(status_code=404, detail="Cover letter not generated yet for this tone")
    with open(txt_path) as f:
        return {"job_id": job_id, "tone": tone, "text": f.read().strip()}


def _cl_txt_path(job_id: str, tone: str) -> str:
    return os.path.join(OUTPUT_DIR, job_id, f"cover_letter_{tone}.txt")


def _cl_html(text: str, tone: str = "", include_meta: bool = True) -> str:
    def _block(b):
        lines = b.strip().splitlines()
        return f"<p>{lines[0]}</p>" if len(lines) == 1 else "<p>" + "<br>".join(lines) + "</p>"
    paragraphs = "".join(_block(b) for b in text.split("\n\n") if b.strip())
    meta = f'<div class="meta">Tone: {tone}</div>' if include_meta else ""
    return f"""<!DOCTYPE html><html><head><meta charset="UTF-8">
<style>
  body {{ font-family: Georgia, serif; max-width: 680px; margin: 60px auto; padding: 0 24px;
          color: #1a1a1a; line-height: 1.7; font-size: 1.05rem; }}
  p {{ margin-bottom: 1.2em; }}
  .meta {{ font-size: 0.78rem; color: #888; margin-bottom: 2em; border-bottom: 1px solid #eee; padding-bottom: 0.8em; }}
</style></head><body>
{meta}
{paragraphs}
</body></html>"""


@router.get("/{job_id}/cover_letter/html")
def get_cover_letter_html(job_id: str, tone: str = "professional"):
    from fastapi.responses import HTMLResponse
    tone = tone.lower()
    txt_path = _cl_txt_path(job_id, tone)
    if not os.path.exists(txt_path):
        raise HTTPException(status_code=404, detail="Cover letter not generated yet for this tone")
    with open(txt_path) as f:
        text = f.read().strip()
    return HTMLResponse(content=_cl_html(text, include_meta=False))


@router.get("/{job_id}/cover_letter/pdf")
def get_cover_letter_pdf(job_id: str, tone: str = "professional"):
    from fastapi.responses import FileResponse
    from app.core.pdf_export import export_pdf
    tone = tone.lower()
    txt_path = _cl_txt_path(job_id, tone)
    if not os.path.exists(txt_path):
        raise HTTPException(status_code=404, detail="Cover letter not generated yet for this tone")
    with open(txt_path) as f:
        text = f.read().strip()
    pdf_path = os.path.join(OUTPUT_DIR, job_id, f"cover_letter_{tone}.pdf")
    export_pdf(_cl_html(text, include_meta=False), pdf_path)
    return FileResponse(
        pdf_path,
        media_type="application/pdf",
        filename=f"cover_letter_{tone}.pdf",
    )


@router.get("/{job_id}/cover_letter/docx")
def get_cover_letter_docx(job_id: str, tone: str = "professional"):
    from fastapi.responses import FileResponse
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    tone = tone.lower()
    txt_path = _cl_txt_path(job_id, tone)
    if not os.path.exists(txt_path):
        raise HTTPException(status_code=404, detail="Cover letter not generated yet for this tone")
    with open(txt_path) as f:
        text = f.read().strip()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.15)
        section.right_margin  = Inches(1.15)

    # Default paragraph style
    style = doc.styles["Normal"]
    font  = style.font
    font.name = "Georgia"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(10)

    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(block.replace("\n", " "))
        run.font.name = "Georgia"
        run.font.size = Pt(11)

    docx_path = os.path.join(OUTPUT_DIR, job_id, f"cover_letter_{tone}.docx")
    doc.save(docx_path)
    return FileResponse(
        docx_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"cover_letter_{tone}.docx",
    )


@router.get("/{job_id}/cover_letter/combined_pdf")
def get_combined_pdf(job_id: str, tone: str = "professional", db: Session = Depends(get_db)):
    from fastapi.responses import FileResponse
    from pypdf import PdfWriter, PdfReader
    from app.core.pdf_export import export_pdf

    tone = tone.lower()

    # Cover letter PDF — generate fresh from txt so edits are always reflected
    txt_path = _cl_txt_path(job_id, tone)
    if not os.path.exists(txt_path):
        raise HTTPException(status_code=404, detail="Cover letter not generated yet for this tone")
    with open(txt_path) as f:
        cl_text = f.read().strip()
    cl_pdf_path = os.path.join(OUTPUT_DIR, job_id, f"cover_letter_{tone}.pdf")
    export_pdf(_cl_html(cl_text, include_meta=False), cl_pdf_path)

    variant = (
        db.query(Variant)
        .filter(Variant.job_id == job_id)
        .order_by(Variant.created_at.desc())
        .first()
    )
    if not variant:
        raise HTTPException(status_code=404, detail="No generated variant found for this job")
    resume_pdf_path = os.path.join(variant.output_path, "generated", "resume.pdf")
    if not os.path.exists(resume_pdf_path):
        raise HTTPException(status_code=404, detail="Resume PDF not found — regenerate the variant")

    # Merge: cover letter first, resume second
    writer = PdfWriter()
    for path in (cl_pdf_path, resume_pdf_path):
        reader = PdfReader(path)
        for page in reader.pages:
            writer.add_page(page)

    combined_path = os.path.join(OUTPUT_DIR, job_id, f"combined_{tone}.pdf")
    with open(combined_path, "wb") as f:
        writer.write(f)

    from fastapi.responses import Response
    with open(combined_path, "rb") as fh:
        data = fh.read()
    return Response(
        content=data,
        media_type="application/pdf",
        headers={"Content-Disposition": "inline"},
    )
